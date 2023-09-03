import docx

file = docx.Document('example.docx')

answers = input('Введите ответы через точку с запятой')

answer_list = answers.split(';')

for i in file.paragraphs:
    i.add_run(answer_list)

file.save('result.docx')
print('The document is generated')